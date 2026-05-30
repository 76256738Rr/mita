"""Genera el manual de usuario MITA en formato Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / 'docs' / 'manual-usuario' / 'MANUAL_USUARIO_MITA.docx'
IMG = BASE / 'docs' / 'manual-usuario' / 'pantallas'


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(26, 54, 93)
    return h


def add_figure(doc, image_name, caption):
    path = IMG / image_name
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(113, 128, 150)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('Manual de Usuario', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('Plataforma MITA\nMulti-Intercultural-Transdisciplina Analógica')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        'Gobierno del Estado de México — SEDESA\n'
        'Versión 1.1 · Mayo 2026\n'
        'Proyecto de referencia: EDOMEX DIABETES 2026–2032'
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, '1. Introducción')
    doc.add_paragraph(
        'La Plataforma MITA integra bases interdisciplinarias, análisis geoespacial, evaluación '
        'analógica, perfiles SNI, emisión de dictámenes y una ventanilla ciudadana Edomex en un '
        'flujo de 7 pasos institucional más portal público para reportes ciudadanos.'
    )
    doc.add_paragraph(
        'Contraseña demo de todos los usuarios: mita2026. Reporte ciudadano demo: MITA-CIU-DEMO-00001.'
    )

    add_heading(doc, '2. Acceso al sistema', 1)
    doc.add_paragraph(
        'Abra la URL de la plataforma. En el entorno demo, seleccione una tarjeta de rol en '
        '/login/. Tras el login, el sistema redirige según el perfil (ventanilla, bandeja, expediente).'
    )
    add_figure(doc, '01-login.png', 'Figura 1 — Inicio de sesión por rol')

    add_heading(doc, '3. Roles, permisos y menú', 1)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Usuario', 'Rol', 'Responsabilidad', 'Pantalla inicial']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ('ciudadano', 'Público general', 'Ventanilla Edomex', '/ciudadania/'),
        ('dependencia', 'Dependencia Gubernamental', 'Paso 1 — Registro', 'Nuevo expediente'),
        ('analista', 'Analista MITA', 'Pasos 2–5 — Análisis', 'Bandeja / expediente demo'),
        ('investigador', 'Investigador SNI', 'Paso 6 — Selección', 'Expediente demo paso 6'),
        ('admin', 'Administrador', 'Paso 7 — Dictamen', 'Expediente demo paso 7'),
        ('auditor', 'Auditor', 'Consulta (solo lectura)', 'Expediente dictaminado'),
    ]
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val
    doc.add_paragraph(
        'El menú lateral muestra solo las opciones permitidas para cada rol. '
        'El ciudadano no ve módulos institucionales; los roles institucionales no ven la ventanilla.'
    )

    add_heading(doc, '4. Dashboard y bandeja', 1)
    add_figure(doc, '02-dashboard.png', 'Figura 2 — Panel de control')
    add_figure(doc, '03-bandeja.png', 'Figura 3 — Bandeja de trabajo')

    add_heading(doc, '5. Proceso operativo', 1)
    doc.add_paragraph(
        'Flujo MITA de 7 pasos: Definición → Recopilación → Organización → Analogía → '
        'Opciones → Evaluación SNI → Dictamen. Cada paso tiene un rol responsable.'
    )
    add_figure(doc, '08-proceso-operativo.png', 'Figura 4 — Vista kanban del proceso operativo')

    add_heading(doc, '6. Expedientes', 1)
    add_figure(doc, '04-expedientes.png', 'Figura 5 — Listado de expedientes')
    add_figure(doc, '05-expediente-nuevo.png', 'Figura 6 — Nuevo expediente (Paso 1)')
    add_figure(doc, '06-expediente-paso6.png', 'Figura 7 — Expediente demo en Paso 6')
    add_figure(doc, '07-expediente-dictaminado.png', 'Figura 8 — Expediente dictaminado (auditoría)')

    add_heading(doc, '7. Módulos institucionales', 1)
    add_figure(doc, '09-base-conocimiento.png', 'Figura 9 — Base de Conocimiento')
    add_figure(doc, '10-mapa-geoespacial.png', 'Figura 10 — Mapa geoespacial')
    add_figure(doc, '11-metodologia.png', 'Figura 11 — Metodología MITA')
    add_figure(doc, '12-proyecto-demo.png', 'Figura 12 — Proyecto EDOMEX DIABETES')
    doc.add_paragraph(
        'Motor Analógico (/analogia/): evaluación con criterios A–F y generación de dictamen PDF. '
        'Accesible para admin y analista.'
    )
    add_figure(doc, '13-perfiles-sni.png', 'Figura 13 — Perfiles SNI')

    add_heading(doc, '8. Ventanilla Ciudadana Edomex', 1)
    doc.add_paragraph(
        'Portal exclusivo del usuario ciudadano. Menú: Ventanilla, Registrar problemática, '
        'Ruta gráfica de atención.'
    )
    steps_ciu = [
        'Seleccionar eje temático (Salud, Educación, Empleo, Economía, Seguridad, Agua, Corrupción).',
        'Completar formulario dinámico y municipio Edomex.',
        'En Salud: elegir centro de salud ISSEM/SEDESA de la lista filtrada por municipio.',
        'Usar asistente IA opcional y enviar reporte (folio MITA-CIU-YYYY-NNNNN).',
        'Dar seguimiento: ruta de atención con responsables, fechas y chat.',
    ]
    for i, s in enumerate(steps_ciu, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    add_heading(doc, '9. Dictámenes y reportes', 1)
    add_figure(doc, '14-dictamenes.png', 'Figura 14 — Dictámenes de validación')
    add_figure(doc, '15-reportes.png', 'Figura 15 — Reportes e indicadores')

    add_heading(doc, '10. Recorridos demo sugeridos', 1)
    doc.add_paragraph('Recorrido institucional:', style='List Bullet')
    steps = [
        'Entrar como dependencia y revisar el registro de expediente.',
        'Entrar como analista y explorar MITA-DEMO-2026-ACTIVO y Motor Analógico.',
        'Entrar como investigador, seleccionar opción y avanzar al Paso 7.',
        'Entrar como admin y emitir el Dictamen de Validación MITA.',
        'Entrar como auditor y consultar MITA-DEMO-2026-DICTAMINADO y reportes.',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')
    doc.add_paragraph('Recorrido ciudadano: portal, reporte MITA-CIU-DEMO-00001, chat y ruta gráfica.')

    add_heading(doc, '11. Preguntas frecuentes', 1)
    faqs = [
        ('¿Por qué no veo todo el menú?', 'El menú es personalizado por rol.'),
        ('¿No puedo editar un expediente?', 'Verifique su rol. El auditor solo consulta.'),
        ('¿Cómo avanzo de paso?', 'Complete los campos, guarde y pulse Avanzar al siguiente paso.'),
        ('¿Cómo registro como ciudadano?', 'Login ciudadano → Registrar problemática → elegir eje.'),
        ('¿Cómo reinicio el demo?', 'Ejecute: python manage.py seed_mita'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(q + ' ').bold = True
        p.add_run(a)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f'Manual Word generado: {OUT}')


if __name__ == '__main__':
    main()
